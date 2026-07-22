"""Divan Services — Export service.

Kararları Markdown, JSON veya DOCX formatında dışa aktarır.
"""

from __future__ import annotations

import io
import json
from typing import Any

import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from ..core.enums import ExportFormat
from ..core.interfaces import IExporter
from ..core.models import Decision
from ..core.exceptions import DivanError


class ExportService(IExporter):
    """Çoklu formatta karar dışa aktarma servisi."""

    async def export(self, decision: Decision, fmt: ExportFormat) -> bytes:
        """Kararı belirtilen formata dönüştür."""
        if fmt == ExportFormat.MARKDOWN:
            return self._export_markdown(decision)
        elif fmt == ExportFormat.JSON:
            return self._export_json(decision)
        elif fmt == ExportFormat.DOCX:
            return self._export_docx(decision)
        elif fmt == ExportFormat.PDF:
            raise DivanError("PDF dışa aktarma henüz desteklenmiyor.")
        else:
            raise DivanError(f"Desteklenmeyen format: {fmt}")

    def _export_markdown(self, decision: Decision) -> bytes:
        """Markdown dışa aktar."""
        content = decision.markdown_content or "*İçerik bulunamadı.*"
        md = f"# {decision.reference}\n\n"
        
        if decision.summary:
            md += f"**Özet:** {decision.summary}\n\n"
            
        md += f"**Mahkeme:** {decision.chamber_name or decision.court_type.name}\n"
        
        if decision.esas_no:
            md += f"**Esas No:** {decision.esas_no}\n"
        if decision.karar_no:
            md += f"**Karar No:** {decision.karar_no}\n"
        if decision.decision_date_str:
            md += f"**Tarih:** {decision.decision_date_str}\n"
            
        md += f"\n---\n\n{content}"
        return md.encode("utf-8")

    def _export_json(self, decision: Decision) -> bytes:
        """JSON dışa aktar."""
        return decision.model_dump_json(indent=2).encode("utf-8")

    def _export_docx(self, decision: Decision) -> bytes:
        """DOCX dışa aktar."""
        doc = docx.Document()
        
        # Style configuration
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Başlık
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(decision.reference)
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Metadata tablosu
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        row_cells = table.rows[0].cells
        row_cells[0].text = 'Mahkeme / Daire:'
        row_cells[1].text = decision.chamber_name or decision.court_type.name
        
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Esas Numarası:'
        row_cells[1].text = decision.esas_no or "-"
        
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Karar Numarası:'
        row_cells[1].text = decision.karar_no or "-"
        
        row_cells = table.rows[3].cells
        row_cells[0].text = 'Karar Tarihi:'
        row_cells[1].text = decision.decision_date_str or "-"
        
        # Boşluk
        doc.add_paragraph()
        
        # Özet (varsa)
        if decision.summary:
            doc.add_heading('Özet', level=2)
            doc.add_paragraph(decision.summary)
        
        # İçerik Başlığı
        doc.add_heading('Karar Metni', level=2)
        
        # Basit Markdown parser for DOCX
        # Not: Tam teşekküllü bir MD -> DOCX converter değil, sadece paragrafları ayırır
        # ve basit bold/italic vb. işaretlerini temizler. Gelişmiş dönüşüm için pypandoc
        # kullanılabilir ancak bağımlılıkları ağırlaştırır.
        content = decision.markdown_content or "İçerik bulunamadı."
        paragraphs = content.split('\n\n')
        
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
                
            # Basit temizleme
            p = p.replace('**', '')
            p = p.replace('*', '')
            p = p.replace('##', '')
            p = p.replace('#', '')
            
            doc.add_paragraph(p)
            
        # Byte buffer'a yaz
        stream = io.BytesIO()
        doc.save(stream)
        return stream.getvalue()
